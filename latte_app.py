"""
라떼는 말이야 ☕ — SNU MBA 후배를 위한 AI 선배 챗봇
Streamlit 웹앱 버전
"""

import os, re, json, time, shutil, zipfile
from pathlib import Path
from datetime import datetime
import streamlit as st
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(page_title="라떼는 말이야 ☕", page_icon="☕", layout="wide")

DEFAULT_SESSION_STATE = {
    'messages': [],
    'conversation_history': [],
    'is_first_turn': True,
    'review_mode': False,
}

for key, default in DEFAULT_SESSION_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = default

st.markdown("""
<style>
.stApp { background-color: #FDF6EC; }
[data-testid="stSidebar"] { background-color: #4A2C2A; }
[data-testid="stSidebar"] * { color: #FDF6EC !important; }
[data-testid="stSidebar"] .stButton button {
    color: #4A2C2A !important;
    background-color: #FDF6EC !important;
    border: none;
}
[data-testid="stSidebar"] button {
    background-color: #FDF6EC !important;
    color: #4A2C2A !important;
}
[data-testid="stSidebar"] button p {
    color: #4A2C2A !important;
}
h1 { color: #4A2C2A; }
</style>
""", unsafe_allow_html=True)

_exam_questions = None
@st.cache_resource(show_spinner="☕ 데이터 불러오는 중... (최초 실행 시 시간이 걸려요)")
def load_all():
    import gdown
    import pandas as pd
    from pypdf import PdfReader
    from docx import Document as DocxDocument
    from langchain_core.documents import Document as LCDocument
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_openai import OpenAIEmbeddings
    from langchain_chroma import Chroma
    from langchain.chat_models import init_chat_model
    from langchain.messages import HumanMessage as _HM
    from langchain.agents import create_agent
    from langchain.tools import tool
    from langchain_tavily import TavilySearch
    from langchain_core.messages import ToolMessage
    import gspread
    from google.oauth2.service_account import Credentials

    import shutil
    import zipfile
    from pathlib import Path
    import gdown

    # 현재 Google Drive의 Latte 폴더 ID
    DRIVE_FOLDER_ID = '1EcjtGxTWVPqrL0-2x0Dz4aualr_dD89P'
    DATA_DIR = 'Latte'

    # 1. 예전에 받아둔 폴더가 남아있으면 지우고 새로 받는다
    shutil.rmtree(DATA_DIR, ignore_errors=True)

    # 2. Google Drive에서 Latte 폴더 다운로드
    gdown.download_folder(
        id=DRIVE_FOLDER_ID,
        output=DATA_DIR,
        quiet=False,
        use_cookies=False
    )

    data_dir = Path(DATA_DIR)

    # 3. 다운로드된 파일 확인
    print("다운로드된 파일:")
    for p in data_dir.rglob("*"):
        print(" -", p)

    # 4. Latte 폴더 안에서 zip 파일 찾기
    zip_candidates = list(data_dir.rglob("*.zip"))
    assert zip_candidates, "Latte 폴더에서 zip 파일을 찾지 못했습니다."

    # zip 파일이 하나라고 가정
    SYLLABUS_ZIP = zip_candidates[0]
    print(f"\nzip 파일: {SYLLABUS_ZIP.name}")

    # 5. zip 압축 풀기
    syllabus_extract_dir = data_dir / "syllabus_extracted"
    shutil.rmtree(syllabus_extract_dir, ignore_errors=True)
    syllabus_extract_dir.mkdir(exist_ok=True)

    with zipfile.ZipFile(SYLLABUS_ZIP, "r") as zip_ref:
        zip_ref.extractall(syllabus_extract_dir)

    print("zip 압축 해제 완료")

    # 6. 압축 푼 폴더에서 PDF/DOCX 찾기
    SYLLABUS_FILES = sorted(
        list(syllabus_extract_dir.rglob("*.pdf")) +
        list(syllabus_extract_dir.rglob("*.docx"))
    )

    # 7. Latte 폴더 안에서 엑셀 파일 찾기
    review_xlsx_candidates = [
        p for p in data_dir.rglob("*.xlsx")
        if not p.name.startswith("~$")
    ]

    assert SYLLABUS_FILES, "압축 해제된 zip 파일에서 PDF/DOCX를 찾지 못했습니다."
    assert review_xlsx_candidates, "Latte 폴더에서 엑셀 파일을 찾지 못했습니다."

    REVIEW_XLSX = review_xlsx_candidates[0]

    print(f"\n강의계획서 파일 {len(SYLLABUS_FILES)}개 발견:")
    for p in SYLLABUS_FILES:
        print(" -", p.name)

    print(f"\n엑셀 파일: {REVIEW_XLSX.name}")

    import pandas as pd

    # 한글 시트
    raw_kr = pd.read_excel(REVIEW_XLSX, sheet_name='Data DB', header=1)
    raw_kr.columns = [str(c).strip() for c in raw_kr.columns]

    # 영어 시트
    raw_en = pd.read_excel(REVIEW_XLSX, sheet_name='Data DB_Eng', header=1)
    raw_en.columns = [str(c).strip() for c in raw_en.columns]
    raw_en = raw_en.rename(columns={
        'Student': '수강자',
        'Course Title': '과목명',
        'Professor': '교수',
        'Language': '진행 언어',
        'Term/Module': '수강 시기(모듈)',
        'Exam (O/X)': '시험 (O/X)',
        'Team Project (O/X)': '팀플 (O/X)',
        'Individual Assignment (O/X)': '개인과제 (O/X)',
        'Comments / Review': '후기(자유 서술)',
    })

    # 합치기
    raw = pd.concat([raw_kr, raw_en], ignore_index=True)
    raw.columns = [str(c).strip() for c in raw.columns]

    reviews_df = raw[['수강자', '과목명', '교수', '진행 언어', '수강 시기(모듈)', '후기(자유 서술)']].copy()
    reviews_df = reviews_df.dropna(subset=['과목명', '후기(자유 서술)'])
    reviews_df = reviews_df.reset_index(drop=True)
    print(f'불러온 후기 개수: {len(reviews_df)}')

    def load_professor_aliases(xlsx_path) -> dict:
        """과목 DB(Q~U열)에서 '교수 매핑(다른 표기)' 정보를 읽어, {원래 교수명: 다른 표기들} 딕셔너리로 만든다.
        과목 DB 영역은 메인 후기 영역과 헤더 행이 달라서(3행), 별도로 읽는다."""
        course_db = pd.read_excel(xlsx_path, sheet_name='Data DB', header=2, usecols='Q:U')
        course_db.columns = [str(c).strip() for c in course_db.columns]

        aliases = {}
        for _, row in course_db.iterrows():
            prof, alt = row['교수'], row['교수.1']
            if pd.notna(prof) and pd.notna(alt):
                aliases[str(prof).strip()] = str(alt).strip()
        return aliases

    professor_aliases = load_professor_aliases(REVIEW_XLSX)
    print(f'교수 별칭 매핑 {len(professor_aliases)}건 로드 완료')


    # ============================================================
    # [추가 v6.1.2] 한국어 <-> 영어 과목명 양방향 매핑 테이블
    # Term 1 Module 1 ~ Term 2 Module 2 중 실제 같은 과목의 한/영 분반 9쌍
    # ============================================================
    COURSE_MAP_PAIRS = [
        # Term 1 Module 1
        ('관리경제학',       'Managerial Economics'),
        ('재무회계',         'Financial Accounting'),
        # Term 1 Module 2
        ('인사 및 조직관리', 'Managing People in Organizations'),
        ('생산서비스운영',   'Operations Management'),
        # Term 1 Module 3
        ('마케팅',           'Marketing'),
        ('경영정보',         'Information Technology'),
        # Term 2 Module 1
        ('전략',             'Strategy'),
        ('재무관리',         'Financial Management'),
        # Term 2 Module 2
        ('조직행동론',       'Organizational Behavior'),
    ]

    course_map = {}
    for kr, en in COURSE_MAP_PAIRS:
        course_map[kr] = en
        course_map[en] = kr

    print(f'한/영 과목명 양방향 매핑 {len(course_map) // 2}쌍 ({len(course_map)}건) 생성:')
    for kr, en in COURSE_MAP_PAIRS:
        print(f'  {kr} <-> {en}')


    from langchain_core.documents import Document as LCDocument

    def build_review_documents(df: pd.DataFrame, aliases: dict) -> list[LCDocument]:
        """후기 데이터프레임 -> Document 리스트로 변환.
        같은 과목에 후기가 여러 개면, 각 문서의 메타데이터에 '해당 과목 총 후기 수'를 같이 기록해서
        답변할 때 'N명의 후기 기반'으로 표시할 수 있게 한다.
        """
        review_count_by_course = df['과목명'].value_counts().to_dict()

        docs = []
        for _, row in df.iterrows():
            course = row['과목명']
            professor = row['교수']
            alt_name = aliases.get(str(professor).strip(), '')
            professor_display = f"{professor} ({alt_name})" if alt_name else professor

            content = (
                f"[과목명: {course}] [교수: {professor_display}] "
                f"[언어: {row['진행 언어']}] [수강시기: {row['수강 시기(모듈)']}]\n"
                f"후기: {row['후기(자유 서술)']}"
            )
            docs.append(LCDocument(
                page_content=content,
                metadata={
                    'type': 'review',
                    'course': course,
                    'professor': professor,
                    'language': row['진행 언어'],
                    'term': row['수강 시기(모듈)'],
                    'author': row['수강자'],
                    'review_count_for_course': review_count_by_course.get(course, 1),
                }
            ))
        return docs

    review_docs = build_review_documents(reviews_df, professor_aliases)
    print(f'생성된 후기 문서 수: {len(review_docs)}')
    print('---')
    print(review_docs[0].page_content[:300])
    print(review_docs[0].metadata)

    #모델 초기화
    # ============================================================
    # OpenAI 모델 초기화용 코드
    # - 배포/공유를 위해 API Key를 코드에 직접 저장하지 않음
    # - .env 파일이 있으면 자동 로드
    # - 없으면 실행 중 입력받음
    # ============================================================
    import os
    import getpass

    try:
        from dotenv import load_dotenv
        load_dotenv()
    except ImportError:
        pass

    from langchain.chat_models import init_chat_model

    def setup_openai_api_key():
        if not os.environ.get("OPENAI_API_KEY"):
            os.environ["OPENAI_API_KEY"] = getpass.getpass("OpenAI API Key를 입력하세요: ")
        return os.environ["OPENAI_API_KEY"]

    setup_openai_api_key()

    # 강의계획서 과목명 추출용 모델 (temperature=0 — 항상 동일한 과목명 추출)
    _extraction_model = init_chat_model(
        model="openai:gpt-4.1-mini",
        temperature=0
    )

    # 강의계획서는 길기 때문에 청크(chunk) 단위로 쪼개서 저장한다.
    # 먼저 각 강의계획서의 앞부분을 AI에게 보여줘서, 정식 과목명을 자동으로 추출한다 (파일명에 의존하지 않음).
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain.chat_models import init_chat_model
    from langchain.messages import HumanMessage as _HumanMessage

    _extraction_model = init_chat_model(model='openai:gpt-4o-mini')

    def extract_course_name_from_syllabus(text: str) -> str:
        """강의계획서 텍스트 앞부분을 보고, 정식 과목명만 한 줄로 추출한다."""
        prompt = (
            "다음은 강의계획서(syllabus) 일부분이다. 이 강의의 정식 과목명만 정확히 한 줄로 답하라. "
            "한국어와 영어 이름이 둘 다 있다면 둘 다 적되, 쉼표로 구분해서 적어라. "
            "설명이나 다른 말은 절대 덧붙이지 말고, 과목명만 출력하라.\n\n"
            f"{text[:1500]}"
        )
        response = _extraction_model.invoke([_HumanMessage(prompt)])
        content = response.content
        if isinstance(content, str):
            return content.strip()
        # content가 리스트 형태인 경우, text 블록만 골라낸다
        text_parts = [
            block.get('text', '') for block in content
            if isinstance(block, dict) and block.get('type') == 'text'
        ]
        return ''.join(text_parts).strip()

    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)

    from pypdf import PdfReader
    from docx import Document as DocxDocument

    def load_pdf_text(path) -> str:
        """PDF의 모든 페이지에서 텍스트를 추출해 하나의 문자열로 합친다."""
        reader = PdfReader(path)
        pages = [page.extract_text() or '' for page in reader.pages]
        return '\n'.join(pages)

    def load_docx_text(path) -> str:
        """DOCX의 문단 + 표 내용을 모두 추출해 하나의 문자열로 합친다."""
        doc = DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells_text:
                    parts.append(' | '.join(cells_text))
        return '\n'.join(parts)

    def load_syllabus_file(path) -> str:
        """확장자에 맞는 방식으로 강의계획서 파일을 읽는다."""
        if path.suffix.lower() == '.pdf':
            return load_pdf_text(path)
        elif path.suffix.lower() == '.docx':
            return load_docx_text(path)
        else:
            raise ValueError(f'지원하지 않는 형식: {path.suffix} (PDF 또는 DOCX만 지원)')

    # syllabus 폴더 안의 파일을 전부 읽어서, (파일명, 텍스트) 쌍으로 모아둔다
    syllabus_texts = [(f.stem, load_syllabus_file(f)) for f in SYLLABUS_FILES]

    syllabus_docs = []
    for file_name, text in syllabus_texts:
        course_name = extract_course_name_from_syllabus(text)
        print(f'{file_name} -> 추출된 과목명: {course_name}')

        chunks = splitter.split_text(text)
        for chunk in chunks:
            syllabus_docs.append(LCDocument(
                page_content=chunk,
                metadata={'type': 'syllabus', 'source_file': file_name, 'course': course_name}
            ))

    print(f'강의계획서 청크 수: {len(syllabus_docs)} (강의계획서 파일 {len(syllabus_texts)}개에서 생성)')

    all_docs = syllabus_docs + review_docs
    print(f'전체 문서 수: {len(all_docs)} (강의계획서 {len(syllabus_docs)} + 후기 {len(review_docs)})')

    from pathlib import Path
    import shutil
    import hashlib
    import json
    import time

    from langchain_chroma import Chroma
    from langchain_openai import OpenAIEmbeddings




    embeddings = OpenAIEmbeddings(model='text-embedding-3-small')

    SYLLABUS_CHROMA_DIR = 'chroma_syllabus_db'
    REVIEW_CHROMA_DIR = 'chroma_review_db'
    SYLLABUS_COLLECTION = 'latte_syllabus_db'
    REVIEW_COLLECTION = 'latte_review_db'

    # -----------------------------
    # 1. 강의계획서 벡터DB
    # -----------------------------
    print('강의계획서 DB를 새로 만듭니다.')
    syllabus_vectorstore = Chroma.from_documents(
        documents=syllabus_docs,
        embedding=embeddings,
        collection_name=SYLLABUS_COLLECTION,
        collection_metadata={'hnsw:space': 'cosine'},
    )
    
    # -----------------------------
    # 2. 수강후기 벡터DB
    # -----------------------------
    print('수강후기 DB를 새로 만듭니다.')
    review_vectorstore = Chroma.from_documents(
        documents=review_docs,
        embedding=embeddings,
        collection_name=REVIEW_COLLECTION,
        collection_metadata={'hnsw:space': 'cosine'},
    )
   
    # syllabus + review 두 DB를 하나로 합친 통합 검색 래퍼
    class CombinedVectorStore:
        def __init__(self, *stores):
            self._stores = stores

        def similarity_search(self, query, k=10):
            k_each = max(1, k // len(self._stores))
            results = []
            for store in self._stores:
                results.extend(store.similarity_search(query, k=k_each))
            return results[:k]


    vectorstore = CombinedVectorStore(syllabus_vectorstore, review_vectorstore)
    print('통합 vectorstore 준비 완료')

    # ── Google Sheets 연결 ──
    review_sheet = None
    try:
        SCOPES = [
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/drive'
        ]
        if 'google_sheets' in st.secrets:
            creds = Credentials.from_service_account_info(
                json.loads(st.secrets['google_sheets']['credentials']), scopes=SCOPES)
        elif 'GOOGLE_SHEETS_CREDENTIALS' in st.secrets:
            creds = Credentials.from_service_account_info(
                json.loads(st.secrets['GOOGLE_SHEETS_CREDENTIALS']), scopes=SCOPES)
        else:
            creds = Credentials.from_service_account_file(
                'halogen-chemist-437121-t6-c46e0c6ddb1e.json', scopes=SCOPES)
        gc = gspread.authorize(creds)
        review_sheet = gc.open_by_key('1VZ60mjNmSb-rwde7uYvcF21Se39skhbnH2Tlz-mJ4ew').worksheet('reviews')
    except Exception:
        review_sheet = None

    # 3단계: Google Sheets에 쌓인 신규 후기를 불러와서 벡터DB에 반영
    # 노트북을 새로 실행할 때마다 자동으로 적용 — 다른 사람이 챗봇으로 남긴 후기도 바로 반영됨

    def load_reviews_from_sheets():
        """Sheets의 reviews 탭에서 신규 후기를 불러와 벡터DB와 메모리에 추가한다."""
        try:
            records = review_sheet.get_all_records()
        except Exception as e:
            print(f'Sheets 로드 실패: {e}')
            return

        if not records:
            print('Sheets에 신규 후기 없음')
            return

        added = 0
        for row in records:
            course = str(row.get('과목명', '')).strip()
            professor = str(row.get('교수', '')).strip()
            language = str(row.get('진행 언어', '')).strip()
            term = str(row.get('입력날짜', '')).strip()
            review_text = str(row.get('후기(자유 서술)', '')).strip()
            author = str(row.get('수강자', '신규입력')).strip()

            if not course or not review_text:
                continue

            alt_name = professor_aliases.get(professor, '')
            professor_display = f"{professor} ({alt_name})" if alt_name else professor

            new_doc = LCDocument(
                page_content=(
                    f'[과목명: {course}] [교수: {professor_display}] '
                    f'[언어: {language}] [수강시기: {term}]\n'
                    f'후기: {review_text}'
                ),
                metadata={
                    'type': 'review', 'course': course, 'professor': professor,
                    'language': language, 'term': term, 'author': author,
                    'review_count_for_course': 0,
                }
            )
            review_vectorstore.add_documents([new_doc])
            all_docs.append(new_doc)
            review_docs.append(new_doc)
            added += 1

        print(f'✅ Sheets에서 {added}건의 신규 후기를 불러와 반영했어요.')

    load_reviews_from_sheets()

    def format_sources(docs: list[LCDocument]) -> str:
        """검색된 문서를 LLM에게 보여줄 텍스트로 정리한다."""
        return '\n\n'.join(f'Source: {d.metadata}\nContent: {d.page_content}' for d in docs)

    def n_reviews_from(docs: list[LCDocument]) -> int:
        """검색된 문서들 중 'review' 타입의 과목명을 파악해서, all_docs 전체에서 해당 과목 후기 수를 센다."""
        review_docs = [d for d in docs if d.metadata.get('type') == 'review']
        if not review_docs:
            return 0
        # 검색된 review 문서들의 과목명 목록
        courses = {d.metadata.get('course', '') for d in review_docs}
        # all_docs 전체에서 해당 과목 후기 수 합산
        return sum(
            1 for d in all_docs
            if d.metadata.get('type') == 'review' and d.metadata.get('course', '') in courses
        )

    from langchain.agents import create_agent
    from langchain.tools import tool
    from langchain.chat_models import init_chat_model

    # 기출문제 원문 임시 저장 (Python이 직접 붙여서 LLM이 요약 못 하게)

    def normalize_name(text: str) -> set[str]:
        """이름 비교를 위해 쉼표/하이픈/슬래시를 다 공백으로 바꾸고 단어 단위로 쪼갠다.
        'Jeong-Yeon Lee'와 'Jeong Yeon Lee'처럼 표기가 다른 경우도 같게 취급하기 위함."""
        cleaned = text.replace(',', ' ').replace('-', ' ').replace('/', ' ')
        return set(cleaned.split())


    # --- 검색추출기(retriever) 도구 정의 (의미 기반 검색 + 토큰/정규화 매칭)
    @tool(response_format='content_and_artifact')
    def retriever_tool(query: str) -> tuple[str, list[LCDocument]]:
        """강의계획서와 강의 후기에서 질문과 관련된 내용을 검색한다."""
        global _exam_questions
        
        retrieved_docs = vectorstore.similarity_search(query, k=10)
        query_lower = query.lower()

        # 공백/하이픈 제거 정규화 — Lee Jeongyeon, lee jeongyeon 등 다양한 표기 대응
        def norm(text):
            return text.lower().replace(' ', '').replace('-', '')

        review_docs_only = [d for d in all_docs if d.metadata.get('type') == 'review']

        # 1. 과목명 매칭 — 토큰 포함 검색 (긴 과목명 우선)
        # '임재현 생산' → ['임재현', '생산'] 토큰이 과목명+교수명 텍스트에 모두 포함되면 매칭
        known_courses = sorted(
            {d.metadata.get('course', '') for d in review_docs_only},
            key=len, reverse=True
        )
        query_tokens = query_lower.split()

        def course_token_match(course):
            # 해당 과목의 교수명도 같이 검색 텍스트에 포함
            profs = {d.metadata.get('professor', '') for d in review_docs_only if d.metadata.get('course') == course}
            search_text = course.lower() + ' ' + ' '.join(p.lower() for p in profs)
            return all(token in search_text for token in query_tokens)

        # 먼저 기존 방식(과목명 전체 포함)으로 시도 → 없으면 토큰 매칭으로 시도
        matched_course = next((c for c in known_courses if c and c.lower() in query_lower), None)
        if not matched_course:
            matched_course = next((c for c in known_courses if c and course_token_match(c)), None)

        # [v6.1.2] 과목명이 잡혔는데 교수명이 다른 분반인 경우 -> 대응 언어 과목으로 전환
        if matched_course:
            matched_course_profs = {
                d.metadata.get('professor', '').lower()
                for d in review_docs_only
                if d.metadata.get('course') == matched_course
            }
            matched_course_prof_norms = set()
            for p in matched_course_profs:
                matched_course_prof_norms.add(norm(p))
                for orig_p, alt_val in professor_aliases.items():
                    if norm(orig_p) == norm(p):
                        for a in alt_val.split(','):
                            matched_course_prof_norms.add(norm(a.strip()))

            all_known_profs = {d.metadata.get('professor', '') for d in review_docs_only}
            query_norm_str = norm(query_lower)
            query_professor = None
            for p in all_known_profs:
                if not p:
                    continue
                alt = professor_aliases.get(p, '')
                alt_parts = [a.strip() for a in alt.split(',') if a.strip()]
                p_norm = norm(p)
                alt_norms = [norm(a) for a in alt_parts]
                if (p.lower() in query_lower or
                    p_norm in query_norm_str or
                    any(a.lower() in query_lower for a in alt_parts) or
                    any(an in query_norm_str for an in alt_norms if len(an) > 2)):
                    query_professor = p
                    break

            if query_professor:
                query_prof_norm = norm(query_professor)
                query_prof_alt = professor_aliases.get(query_professor, '')
                query_prof_alt_norms = {norm(a.strip()) for a in query_prof_alt.split(',') if a.strip()}
                all_query_prof_norms = {query_prof_norm} | query_prof_alt_norms

                prof_mismatch = matched_course_prof_norms.isdisjoint(all_query_prof_norms)

                if prof_mismatch:
                    alt_course = course_map.get(matched_course)
                    if alt_course:
                        alt_course_profs = {
                            d.metadata.get('professor', '').lower()
                            for d in review_docs_only
                            if d.metadata.get('course') == alt_course
                        }
                        alt_course_prof_norms = set()
                        for p in alt_course_profs:
                            alt_course_prof_norms.add(norm(p))
                            for orig_p, alt_val in professor_aliases.items():
                                if norm(orig_p) == norm(p):
                                    for a in alt_val.split(','):
                                        alt_course_prof_norms.add(norm(a.strip()))
                        if not alt_course_prof_norms.isdisjoint(all_query_prof_norms):
                            matched_course = alt_course

        # 2. 교수명 매칭 — 정규화(공백/하이픈 제거)로 다양한 영문 표기 대응
        matched_professor = None
        known_professors = {d.metadata.get('professor', '') for d in review_docs_only}
        for p in known_professors:
            if not p:
                continue
            alt = professor_aliases.get(p, '')
            alt_parts = [a.strip() for a in alt.split(',') if a.strip()]
            # 정규화 비교: leejeongyeon == lee jeongyeon == Lee-Jeongyeon 모두 매칭
            p_norm = norm(p)
            alt_norms = [norm(a) for a in alt_parts]
            query_norm = norm(query_lower)
            if (p.lower() in query_lower or
                p_norm in query_norm or
                any(a.lower() in query_lower for a in alt_parts) or
                any(an in query_norm for an in alt_norms if len(an) > 4)):
                matched_professor = p
                break

        # 3. 필터링 적용
        if matched_course:
            syllabus_docs_filtered = [d for d in retrieved_docs if d.metadata.get('type') != 'review']
            # 교수님이 특정된 경우 과목명 + 교수님 둘 다 필터
            if matched_professor:
                alt = professor_aliases.get(matched_professor, '')
                pnames = {matched_professor.lower()} | {a.lower() for a in alt.split(',') if a.strip()}
                review_docs_filtered = [d for d in all_docs if
                                         d.metadata.get('type') == 'review' and
                                         d.metadata.get('course', '').lower() == matched_course.lower() and
                                         d.metadata.get('professor', '').lower() in pnames]
            else:
                review_docs_filtered = [d for d in all_docs if
                                         d.metadata.get('type') == 'review' and
                                         d.metadata.get('course', '').lower() == matched_course.lower()]
            retrieved_docs = syllabus_docs_filtered + review_docs_filtered
        elif matched_professor:
            alt = professor_aliases.get(matched_professor, '')
            alt_parts = [a.strip() for a in alt.split(',') if a.strip()]
            matched_professor_names = {matched_professor.lower()} | {a.lower() for a in alt_parts}
            syllabus_docs_filtered = [d for d in retrieved_docs if d.metadata.get('type') != 'review']
            review_docs_filtered = [d for d in all_docs if
                                     d.metadata.get('type') == 'review' and
                                     d.metadata.get('professor', '').lower() in matched_professor_names]
            retrieved_docs = syllabus_docs_filtered + review_docs_filtered

       # 기출문제 원문 따로 저장 (Python이 직접 붙여서 LLM이 요약 못 하게)
       # 사람마다 기출문제 앞에 쓰는 표현이 다를 수 있어서 여러 키워드로 시도
        exam_texts = []
        exam_keywords = [
    '기출시험문제 공유',
    '기출시험문제',
    '시험기출문제',
    '기출문제',
    '기출 문제',
    '시험문제',
    'sample exam',
    'practice exam',
    'Sample exam',
    'Practice exam',
]
        for d in retrieved_docs:
            if d.metadata.get('type') == 'review':
                content = d.page_content
                for keyword in EXAM_KEYWORDS:
                    if keyword in content:
                        # 키워드 뒤의 내용(실제 문제)만 잘라냄
                        exam_part = content.split(keyword, 1)[1].strip()
                        # 키워드 바로 뒤에 ':'가 붙어있으면 제거 (예: '기출문제: 1. ...')
                        if exam_part.startswith(':'):
                            exam_part = exam_part[1:].strip()
                        if exam_part:
                            exam_texts.append(exam_part)
                        break  # 키워드 하나 찾으면 다음 키워드는 시도 안 함
        _exam_questions = '\n\n'.join(exam_texts) if exam_texts else None

        return format_sources(retrieved_docs), retrieved_docs


    # --- 과목 DB(엑셀)에서 교수명-과목 전체 매핑을 미리 만들어둔다 (후기 유무와 무관하게 공식 과목 목록을 알 수 있음)
    def build_course_catalog(xlsx_path) -> list[dict]:
        """과목 DB(Q~U열)를 읽어, 과목별로 {과목명, 교수, 별칭} 딕셔너리 리스트를 만든다."""
        course_db = pd.read_excel(xlsx_path, sheet_name='Data DB', header=2, usecols='Q:U')
        course_db.columns = [str(c).strip() for c in course_db.columns]

        catalog = []
        for _, row in course_db.iterrows():
            course, prof, alt = row.get('과목명'), row.get('교수'), row.get('교수.1')
            if pd.notna(course) and pd.notna(prof):
                catalog.append({
                    'course': str(course).strip(),
                    'professor': str(prof).strip(),
                    'alt': str(alt).strip() if pd.notna(alt) else '',
                })
        return catalog

    course_catalog = build_course_catalog(REVIEW_XLSX)


    # --- 교수님 전체 과목 조회 도구 (과목 DB 직접 조회 — 후기 유무와 상관없이 공식 과목 전부를 찾는다)
    @tool(response_format='content_and_artifact')
    def professor_courses_tool(professor_name: str) -> tuple[str, list[LCDocument]]:
        """특정 교수님이 공식적으로 담당하는 모든 과목을 정확하게 찾는다 (후기가 없는 과목도 포함된다).
        '이 교수님 수업/과목 다 뭐 있어?'처럼 한 교수님의 전체 과목을 물을 때 사용한다."""
        query_words = normalize_name(professor_name)

        def norm(text):
            # 공백/하이픈/쉼표 제거 후 소문자
            return text.lower().replace(' ', '').replace('-', '').replace(',', '')

        def norm_sorted(text):
            # 알파벳만 추출해서 정렬 — 순서가 달라도 같은 이름으로 인식
            # leejeongyeon / jeongyeonlee / LeeJeongYeon 모두 같아짐
            letters = sorted(norm(text))
            return ''.join(letters)

        matched_courses = []
        query_norm_sorted = norm_sorted(professor_name)

        for entry in course_catalog:
            candidate_text = f"{entry['professor']} {entry['alt']}"
            candidate_words = normalize_name(candidate_text)

            # 방식 1: 단어 단위 매칭 (Lee Jeong Yeon 등 띄어쓴 경우)
            word_match = query_words and query_words.issubset(candidate_words)

            # 방식 2: 정렬 비교 — 순서/붙임/하이픈 달라도 같은 이름으로 인식
            # 교수명과 별칭을 쉼표/슬래시로 쪼개서 각각 비교 (교수 2명인 경우 대응)
            candidate_names = []
            for raw in [entry['professor'], entry['alt']]:
                for sep in ['/', ',']:
                    parts = [p.strip() for p in raw.split(sep) if p.strip()]
                    candidate_names.extend(parts)
            candidate_names = list(set(candidate_names))  # 중복 제거
            sort_match = any(norm_sorted(n) == query_norm_sorted for n in candidate_names)

            if word_match or sort_match:
                matched_courses.append(entry)

        # 매칭된 과목이 없으면 확인 요청
        if not matched_courses:
            return (
                f'입력하신 "{professor_name}" 교수님 성함을 찾지 못했어요. '
                f'성함을 다시 한 번 확인해 주시겠어요? '
                f'비슷한 이름이 있더라도 정확한 성함으로만 검색돼요. '
                f'(예: 이재호 → 이제호처럼 한 글자 차이도 다른 분으로 인식해요)'
            ), []

        result_docs = []
        summary_lines = []
        for entry in matched_courses:
            course = entry['course']
            course_reviews = [
                d for d in all_docs
                if d.metadata.get('type') == 'review' and d.metadata.get('course') == course
            ]
            if course_reviews:
                result_docs.extend(course_reviews)
                summary_lines.append(f"- {course}: 후기 {len(course_reviews)}건 있음")
            else:
                summary_lines.append(f"- {course}: 아직 후기 없음 (공식 과목 목록에는 있음)")

        summary = f"{professor_name} 교수님 담당 과목 목록:\n" + '\n'.join(summary_lines)
        return summary + '\n\n' + format_sources(result_docs), result_docs

    # --- 일반 지식 답변 도구
    @tool
    def general_knowledge_tool(question: str) -> str:
        """내부 자료(강의계획서/후기)에 관련 정보가 전혀 없거나 매우 부족할 때만 사용한다.
        MBA 수업/커리어 일반론에 대한 합리적인 일반 지식을 바탕으로 답한다.
        이 도구를 호출하면 '내부 자료 기반'이 아니라 'AI 일반 답변'이라는 점이 사용자에게 표시된다."""
        return (
            "이 질문은 내부 자료(강의계획서/후기)만으로는 답하기 어렵다. "
            "MBA 커리큘럼과 직무 전반에 대한 일반적인 지식을 바탕으로, "
            "신중하고 균형 잡힌 답변을 만들어라. 확실하지 않은 정보를 단정적으로 말하지 않는다."
        )


    # --- 웹 검색 도구
    from langchain_tavily import TavilySearch

    web_search_tool = TavilySearch(
        max_results=3,
        name='web_search_tool',
        description=(
            "회사명, 기관명, 또는 특정 직무가 정확히 무엇을 하는 곳/일인지 파악하기 위해 웹을 검색한다. "
            "진로 추천 질문에서, 사용자가 언급한 회사/기관/직무가 생소하거나 구체적인 업무 내용을 알아야 할 때 가장 먼저 사용한다."
        ),
    )


    # --- 진로 맞춤 과목 추천 도구
    @tool(response_format='content_and_artifact')
    def career_recommendation_tool(keywords: str) -> tuple[str, list[LCDocument]]:
        """직무/회사 관련 키워드로 강의계획서와 후기에서 관련성 높은 내용을 검색한다.
        회사/기관명을 그대로 넣기보다는, web_search_tool로 파악한 업무 특성 키워드를 넣는 것이 더 정확하다."""
        retrieved_docs = vectorstore.similarity_search(keywords, k=15)
        return format_sources(retrieved_docs), retrieved_docs

    # --- 과목 비교 도구
    @tool(response_format='content_and_artifact')
    def course_comparison_tool(query: str) -> tuple[str, list[LCDocument]]:
        """두 개 이상의 과목을 비교하기 위해 각 과목의 강의계획서와 후기를 검색한다.
        사용자가 'A랑 B 비교해줘', 'A vs B', 'A와 B 중 뭐 들을까?',
        'Compare A and B'처럼 물을 때 사용한다.
        """

        import re
        from difflib import get_close_matches

        # 전체 문서에서 알려진 과목명 수집
        known_courses = sorted(
            {d.metadata.get('course', '') for d in all_docs if d.metadata.get('course')},
            key=len,
            reverse=True
        )

        query_lower = query.lower()
        matched_courses = []

        # 1. 질문 안에 과목명이 그대로 포함된 경우
        for course in known_courses:
            if course and course.lower() in query_lower:
                matched_courses.append(course)

        # 2. 그대로 포함되지 않으면 구분어 기준으로 쪼개서 fuzzy match
        if len(matched_courses) < 2:
            cleaned = query

            split_tokens = [
                '비교해줘', '비교해 주세요', '비교', '뭐 들을까', '뭐가 나아',
                '중에', '중에서', '랑', '이랑', '하고', '과', '와',
                'compare', 'Compare', 'vs', 'VS', 'versus', 'between', 'and'
            ]

            for token in split_tokens:
                cleaned = cleaned.replace(token, '|')

            candidates = [part.strip() for part in cleaned.split('|') if len(part.strip()) >= 2]

            for cand in candidates:
                match = get_close_matches(cand, known_courses, n=1, cutoff=0.45)
                if match:
                    matched_courses.append(match[0])

        # 중복 제거
        matched_courses = list(dict.fromkeys(matched_courses))

        # 3. 과목이 2개 이상 잡히면 각 과목의 모든 자료 수집
        comparison_docs = []

        if len(matched_courses) >= 2:
            for course in matched_courses[:3]:  # 최대 3개까지만 비교
                course_docs = [
                    d for d in all_docs
                    if d.metadata.get('course', '').lower() == course.lower()
                ]
                comparison_docs.extend(course_docs)

        # 4. 그래도 못 찾으면 벡터 검색으로 보완
        if not comparison_docs:
            comparison_docs = vectorstore.similarity_search(query, k=12)

        if not comparison_docs:
            return (
                "비교할 과목 정보를 찾지 못했어요. 과목명을 조금 더 정확히 입력해 주세요.\n"
                "예: Marketing Analytics랑 Data-Driven Decision-Making 비교해줘\n"
                "Example: Compare Marketing Analytics and Data-Driven Decision-Making."
            ), []

        return format_sources(comparison_docs), comparison_docs

    # --- 후기 개수 집계 도구
    @tool
    def count_reviews_tool(course_name: str = '') -> str:
        """수강후기 개수를 과목별로 집계한다. course_name이 비어있으면 전체 과목을 후기 많은 순으로 보여준다."""
        from collections import Counter
        counts = Counter(d.metadata.get('course', '') for d in review_docs)
        if course_name:
            count = counts.get(course_name, 0)
            return f'{course_name} 후기: {count}건'
        lines = [f'{c}: {n}건' for c, n in counts.most_common()]
        return '과목별 후기 수: ' + ', '.join(lines)


    # --- 시스템 프롬프트
    SYSTEM_PROMPT = '''너는 SNU MBA 후배들에게 과목과 학교생활에 대해 알려주는 AI 선배 챗봇 "라떼는 말이야"이다.

    [기본 동작 — 도구 선택]
    질문이 과목, 교수님, 진로/직무, 학교생활 등 우리 데이터와 관련 있다면 반드시 적절한 도구를 먼저 호출한 뒤 답변한다.
    인사, 감사 표현, 일상적인 잡담처럼 검색이 필요 없는 내용이라면 도구를 호출하지 않고 친근하게 바로 답한다.
    사용자가 특정 과목명을 언급하며 질문한다면 반드시 retriever_tool을 사용한다.
    사용자가 두 개 이상의 과목을 비교해 달라고 하면 반드시 course_comparison_tool을 사용한다.
    예를 들어 "A랑 B 비교해줘", "A vs B", "A와 B 중 뭐 들을까?", "Compare A and B", "Which one should I take between A and B?" 같은 질문은 course_comparison_tool을 사용한다.
    과목 비교 질문에서는 retriever_tool보다 course_comparison_tool을 우선 사용한다.
    과목 비교 답변은 단순 요약이 아니라 수강신청 의사결정에 도움이 되도록 비교 중심으로 작성한다.
    비교 항목은 가능한 경우 "난이도", "과제/팀플 부담", "시험 부담", "실무성", "추천 대상", "주의할 점", "최종 추천"을 포함한다.
    검색 결과에 없는 비교 항목은 지어내지 말고 생략한다.
    사용자가 과목명 없이 교수님 이름만으로 전체 과목 목록을 묻는다면 professor_courses_tool을 사용한다.
    질문에 과목명이 하나라도 포함되어 있다면 professor_courses_tool을 사용하지 않는다.
    professor_courses_tool을 사용했다면, 결과에 나온 모든 과목을 빠짐없이 언급한다. 후기가 없는 과목은 "아직 후기가 없다"고 명시한다.
    같은 교수님이 한국어/영어 과목명으로 각각 다른 수업을 가르치는 경우, 서로 다른 두 과목일 수도 같은 과목의 한/영 분반일 수도 있으니 구분해서 명확히 설명한다.
    "가장 후기가 많은 과목", "후기 몇 개" 같은 통계/집계 질문은 count_reviews_tool을 사용한다.
    이전 대화에서 과목이나 교수님이 언급됐다면, "난이도", "시험", "과제", "추천" 같은 후속 질문도 반드시 retriever_tool을 호출해서 내부 자료로 답변한다. 절대 도구 없이 바로 답하지 않는다.

    [진로 추천 — 회사/직무 질문 처리 순서]
    사용자가 관심 직무나 회사를 언급하며 과목 추천을 묻는다면 다음 순서를 반드시 지킨다.
    1. web_search_tool로 그 회사/기관/직무가 무엇을 하는 곳인지, 어떤 업무 특성을 가지는지 검색한다.
    2. 검색 결과에서 핵심 업무 키워드를 2~4개 뽑는다.
    3. 그 키워드로 career_recommendation_tool을 호출해 내부 강의계획서/후기를 검색한다.
    4. 내부 자료에서 관련 과목을 찾으면 그 내용을 우선 근거로 사용한다.
    5. 내부 자료가 부족하면, 그 사실을 솔직히 밝히고 web_search_tool로 알게 된 직무 특성과 일반 지식을 바탕으로 보완 추천한다.
    추천 과목은 반드시 내부 자료에 실제로 존재하는 SNU MBA 과목명만 언급한다. 내부 자료에 없는 과목명은 절대 만들어내거나 추천하지 않는다.

    [2단계 — 내부 자료가 부족할 때 일반 답변으로 전환]
    retriever_tool 또는 professor_courses_tool을 호출했는데 관련 자료가 0건이면 반드시, 1건이면 질문에 충분히 답하기 어렵다고 판단될 때 general_knowledge_tool을 추가로 호출한 뒤 일반 지식으로 답변한다.
    진로 추천 질문(career_recommendation_tool 사용)에는 이 단계를 적용하지 않는다.
    이 경우 답변 첫 문장에 "이 답변은 우리 학교 내부 자료가 아니라 일반적인 지식을 바탕으로 한 참고용 의견이에요."를 한 번만 포함한다.
    general_knowledge_tool로 답변하는 경우에는 JSON 형식을 사용하지 않는다. 자연스러운 문단형 답변으로 작성한다.
    답변 첫 문장에는 반드시 "이 답변은 우리 학교 내부 자료가 아니라 일반적인 지식을 바탕으로 한 참고용 의견이에요."를 포함한다.

    [정확성 — 가장 중요한 규칙]
    검색 결과에 있는 내용만 근거로 답변한다. 자신의 내부 지식으로 추측하거나 지어내지 않는다.
    단, general_knowledge_tool이나 web_search_tool을 호출한 경우는 예외이며, 이때도 외부 정보임을 명시한다.
    검색 결과에 없는 이름, 수업 내용, 평가 방식은 절대 포함하지 않는다.
    강의계획서와 후기 내용이 다르다면, 실제 경험을 담은 후기를 더 신뢰할 만한 정보로 우선한다.
    여러 학기 후기가 섞여 있다면 최신 정보를 우선하되 "학기/교수님에 따라 다를 수 있다"고 안내한다.
    과목명은 내부 자료에 저장된 표기 그대로만 사용한다.
    영어 과목명을 한국어로 번역하거나, 한국어 과목명을 영어로 번역하지 않는다.
    과목명 옆에 번역이나 설명을 괄호로 추가하지 않는다. 접두어도 절대 붙이지 않는다. 
    과목명을 임의로 변경하거나 유사한 이름으로 바꾸지 않는다. 반드시 내부 자료에 저장된 과목명을 한 글자도 바꾸지 않고 그대로 사용한다.

    [교수님 이름 및 과목명 처리]
    교수님 성함은 검색 결과에 정확히 등장한 표기 그대로만 사용한다. 변형하거나 비슷한 발음으로 추측해 새로운 이름을 만들어내지 않는다. 한글/영문 표기가 달라도 같은 분으로 인식한다.
    사용자가 특정 교수님을 언급한 경우, 그 교수님의 후기와 강의계획서만 답변에 사용한다. 다른 교수님의 후기는 절대 포함하지 않는다.
    사용자가 특정 과목명을 언급한 경우, 그 과목과 일치하는 후기만 사용한다. 대소문자가 달라도 같은 과목으로 인식한다.
    교수님 이름이나 과목명을 검색했는데 자료를 찾지 못한 경우:
    - 우리 학교 과목 목록에 있는 것처럼 보이지만 후기가 없다면 → "아직 이 과목에 대한 후기가 없어요. 곧 채워질 예정이니 조금만 기다려 주세요!"
    - 입력 오류이거나 시스템에 전혀 등록되지 않은 것 같다면 → "입력하신 교수님 성함이나 과목명을 다시 한 번 확인해 주시겠어요?"
    - 판단이 애매하면 전자(아직 후기가 없다는 안내)를 우선한다.

    [내용 순화]
    후기에 욕설, 비속어, 특정인을 비방하는 표현이 있다면 답변에 그대로 옮기지 않고 순화하거나 생략한다. 다만 후기의 핵심 내용은 빠뜨리지 않고 전달한다.

    [답변 형식]
    후기 내용을 아래 JSON 형식으로만 출력한다. 코드 블록(```) 없이 순수한 JSON만 출력한다.
    사용자가 한국어로 질문하면 JSON의 값(value)은 한국어로 작성한다.
    사용자가 영어로 질문하면 JSON의 값(value)은 영어로 작성한다.
    JSON 키(key)는 항상 아래 지정된 한국어 키를 그대로 사용한다.
    후기에 없는 항목은 JSON에 포함하지 않는다(null 포함 금지).
    원문에 없는 내용은 절대 추가하지 않으며, 원문 표현을 최대한 그대로 살린다.
    원문에 쓰인 항목/주제는 그 표현 그대로 언급하고, 없는 개념이나 이론 이름을 추가로 붙이지 않는다.
    검색된 문서에 [과목명: ...] [교수: ...] [언어: ...] [수강시기: ...] 형식의 태그가 있어도 답변에 절대 출력하지 않는다.
    시험기출문제는 별도 파이프라인이 JSON에 자동 삽입하므로 LLM은 해당 항목을 직접 작성하지 않는다.

    검색된 후기가 한국어와 영어로 섞여 있어도, 같은 과목에 대한 후기라면 반드시 하나의 JSON 객체로 통합한다.
    후기 원문의 언어가 달라도 답변 언어는 사용자의 질문 언어를 따른다.
    예를 들어 사용자가 한국어로 질문하면 영어 후기에 나온 내용도 자연스러운 한국어로 번역·요약해서 같은 항목에 합친다.
    사용자가 영어로 질문하면 한국어 후기에 나온 내용도 자연스러운 영어로 번역·요약해서 같은 항목에 합친다.
    같은 과목의 여러 후기에서 같은 항목에 해당하는 내용이 나오면 key를 반복하지 말고 하나의 value 안에 합친다.
    예를 들어 한국어 후기와 영어 후기 모두 수업 진행 방식에 대해 말하고 있다면 "수업방식" key는 한 번만 사용하고, 두 후기의 내용을 하나의 value에 함께 담는다.
    같은 JSON 객체 안에서 동일한 key를 두 번 이상 사용하지 않는다.

    진로 추천 질문이면 아래 형식을 우선 사용한다.
    진로 추천 질문에서는 일반 후기용 키인 "개인과제", "팀프로젝트", "수업분위기", "꿀팁추천대상", "시험", "교수스타일", "기타"를 사용하지 않는다.
    진로 추천 질문에서는 반드시 아래 키만 사용한다:
    "과목명", "추천이유", "얻을수있는역량", "수업방식", "평가방식", "주의할점"
    영어 질문이어도 JSON key는 위의 한국어 키를 그대로 사용하고, value만 영어로 작성한다.

    이때 추천 과목은 반드시 내부 자료에 실제로 존재하는 SNU MBA 과목명만 사용한다.
    JSON의 "과목명" 값은 career_recommendation_tool이 검색해온 문서의 [과목명: ...] 태그에 적힌 이름을 그대로 복사하며, 번역, 수정, 괄호 추가, 접두어 추가를 절대 하지 않는다.
    각 과목마다 사용자의 관심 직무/회사/진로와 연결되는 근거를 "추천이유"에 포함한다.
    후기나 강의계획서에 근거가 없는 항목은 포함하지 않는다.
    내부 자료에서 확인되지 않은 수업 내용, 평가 방식, 역량, 주의사항은 절대 지어내지 않는다.
    영어 질문이면 "추천이유", "얻을수있는역량", "수업방식", "평가방식", "주의할점"의 value를 모두 자연스러운 영어로 작성한다.
    한국어 질문이면 해당 value를 모두 자연스러운 한국어 캐주얼 존댓말로 작성한다.

    진로 추천 과목이 1개이면 단일 JSON 객체로:
    {
      "과목명": "과목명",
      "추천이유": "...",
      "얻을수있는역량": "...",
      "수업방식": "...",
      "평가방식": "...",
      "주의할점": "..."
    }

    진로 추천 과목이 2개 이상이면 반드시 JSON 배열로:
    [
      {
        "과목명": "과목명1",
        "추천이유": "...",
        "얻을수있는역량": "...",
        "수업방식": "...",
        "평가방식": "...",
        "주의할점": "..."
      },
      {
        "과목명": "과목명2",
        "추천이유": "...",
        "얻을수있는역량": "...",
        "수업방식": "...",
        "평가방식": "...",
        "주의할점": "..."
      }
    ]

    일반 과목/교수님 후기 질문이면 아래 형식을 사용한다.

    과목이 1개이면 단일 JSON 객체로:
    과목이 1개이면 후기 언어가 여러 개여도 절대 JSON 배열이나 여러 객체로 나누지 않고, 반드시 단일 JSON 객체 하나로 합쳐서 출력한다.
    같은 과목의 한국어 후기와 영어 후기가 함께 검색되어도 언어별로 항목을 나누지 않는다.
    {
      "수업방식": "...",
      "개인과제": "...",
      "팀프로젝트": "...",
      "수업분위기": "...",
      "꿀팁추천대상": "...",
      "시험": "...",
      "교수스타일": "...",
      "기타": "위 항목에 없는 중요한 내용"
    }

    과목이 2개 이상이면 반드시 JSON 배열로, 각 객체에 "과목명" 키를 포함:
    검색 결과에 여러 과목의 후기가 포함되어 있으면 절대 하나의 종합 요약으로 합치지 않는다.
    반드시 과목별로 분리해서 JSON 배열을 출력한다.
    예를 들어 한 교수님이 여러 과목을 가르친 경우에도, 각 과목을 별도 객체로 나눠 작성한다.
    [
      {
        "과목명": "과목명1",
        "수업방식": "...",
        "시험": "..."
      },
      {
        "과목명": "과목명2",
        "수업방식": "..."
      }
    ]

    관련 없는 다른 과목이나 교수님의 후기 내용은 절대 포함하지 않는다.

    [말투 — 모든 문장에 예외 없이 적용]
    사용자가 한국어로 질문하면 JSON value 안의 텍스트 포함, 모든 문장을 "~해요", "~더라구요", "~하시더라고요", "~한 편이에요" 같은 캐주얼한 존댓말로 작성한다.
    한국어 답변에서는 "~합니다", "~입니다" 같은 격식체와 반말("~이야", "~해", "~같아")을 절대 사용하지 않는다.
    사용자가 영어로 질문하면 JSON value 안의 텍스트 포함, 모든 문장을 자연스럽고 친근한 영어로 작성한다. 한국어 존댓말 어미를 영어 답변에 붙이지 않는다.
    '''

    model = init_chat_model(model='openai:gpt-4.1-mini')

    agent = create_agent(
        model=model,
        tools=[
            retriever_tool,
            professor_courses_tool,
            general_knowledge_tool,
            web_search_tool,
            career_recommendation_tool,
            course_comparison_tool,
            count_reviews_tool,
        ],
        system_prompt=SYSTEM_PROMPT,
    )

    return dict(
        all_docs=all_docs,
        review_docs=review_docs,
        review_vectorstore=review_vectorstore,
        vectorstore=vectorstore,
        professor_aliases=professor_aliases,
        review_sheet=review_sheet,
        agent=agent,
    )


# ── 헬퍼 함수들 (셀 31) ──
import re

TOOL_NAMES = {
    'retriever_tool',
    'professor_courses_tool',
    'general_knowledge_tool',
    'career_recommendation_tool',
    'course_comparison_tool',
    'count_reviews_tool'
}
GENERAL_TOOL_NAME = 'general_knowledge_tool'
CAREER_TOOL_NAME = 'career_recommendation_tool'
WEB_SEARCH_TOOL_NAME = 'web_search_tool'
PROFESSOR_TOOL_NAME = 'professor_courses_tool'

def extract_text(result):
    content = result['messages'][-1].content
    if isinstance(content, str):
        return content
    return '\n'.join(block.get('text', '') for block in content if isinstance(block, dict) and block.get('type') == 'text')

def n_reviews_from(docs):
    review_docs = [d for d in docs if d.metadata.get('type') == 'review']
    if not review_docs:
        return 0
    courses = {d.metadata.get('course', '') for d in review_docs}
    return sum(1 for d in all_docs if d.metadata.get('type') == 'review' and d.metadata.get('course', '') in courses)

def extract_n_reviews(result):
    n = 0
    for message in result['messages']:
        if isinstance(message, ToolMessage) and message.name in TOOL_NAMES and message.name != GENERAL_TOOL_NAME:
            artifact = getattr(message, 'artifact', None)
            if artifact:
                n = n_reviews_from(artifact)
    return n

def extract_json_values_text_for_lang(text):
    """
    언어 감지용 텍스트를 만든다.
    LLM 답변이 JSON이면 한국어 key(수업방식/교수스타일 등)는 제외하고
    value만 모아서 반환한다. JSON이 아니면 원문을 그대로 반환한다.
    """
    try:
        import json as _json_lang
        clean = re.sub(r'```json|```', '', str(text)).strip()
        parsed = _json_lang.loads(clean)
        values = []

        def walk(x):
            if isinstance(x, dict):
                for v in x.values():
                    walk(v)
            elif isinstance(x, list):
                for item in x:
                    walk(item)
            elif isinstance(x, str):
                values.append(x)

        walk(parsed)
        return ' '.join(values) if values else str(text)
    except Exception:
        return str(text)

def is_english(text):
    # JSON 답변은 key가 아니라 value 기준으로 언어를 판단한다.
    content = extract_json_values_text_for_lang(text)
    if len(re.findall(r'[가-힣]', content)) > 0:
        return False
    return len(re.findall(r'[a-zA-Z]', content)) > 0

def collect_searched_pairs(result):
    pairs = []
    seen = set()
    for message in result['messages']:
        if isinstance(message, ToolMessage) and message.name in TOOL_NAMES and message.name != GENERAL_TOOL_NAME:
            artifact = getattr(message, 'artifact', None)
            if not artifact:
                continue
            for d in artifact:
                if d.metadata.get('type') != 'review':
                    continue
                pair = (d.metadata.get('course', ''), d.metadata.get('professor', ''))
                if pair[0] and pair not in seen:
                    seen.add(pair)
                    pairs.append(pair)
    return pairs

def extract_exam_questions_from_result(result):
    """검색된 review 문서에서 시험/기출 관련 원문을 직접 뽑아낸다."""
    exam_texts = []
    exam_keywords = [
        '기출시험문제 공유',
        '기출시험문제',
        '시험기출문제',
        '기출문제',
        '기출 문제',
        '시험문제',
        'sample exam',
        'practice exam',
        'Sample exam',
        'Practice exam',
    ]

    for message in result['messages']:
        if isinstance(message, ToolMessage) and message.name in TOOL_NAMES:
            artifact = getattr(message, 'artifact', None)
            if not artifact:
                continue

            for d in artifact:
                if d.metadata.get('type') != 'review':
                    continue

                content = d.page_content

                for keyword in exam_keywords:
                    if keyword in content:
                        if keyword in ['시험문제', 'sample exam', 'practice exam', 'Sample exam', 'Practice exam']:
                            exam_part = content.strip()
                        else:
                            exam_part = content.split(keyword, 1)[1].strip()

                            if exam_part.startswith(':'):
                                exam_part = exam_part[1:].strip()

                        if exam_part:
                            exam_texts.append(exam_part)

                        break

    return '\n\n'.join(exam_texts)
def get_comparison_sources_text(result):
    """course_comparison_tool이 가져온 문서들을 비교 답변용 텍스트로 정리한다."""
    sources = []

    for message in result['messages']:
        if isinstance(message, ToolMessage) and message.name == 'course_comparison_tool':
            artifact = getattr(message, 'artifact', None)
            if not artifact:
                continue

            for d in artifact:
                course = d.metadata.get('course', '')
                professor = d.metadata.get('professor', '')
                doc_type = d.metadata.get('type', '')
                content = d.page_content

                sources.append(
                    f"[type: {doc_type}] [course: {course}] [professor: {professor}]\n{content}"
                )

    return "\n\n---\n\n".join(sources)

def get_header_info(result, answer):
    is_prof_tool = any(isinstance(m, ToolMessage) and m.name == PROFESSOR_TOOL_NAME for m in result['messages'])

    if is_prof_tool:
        mentioned = []
        seen = set()
        for m in result['messages']:
            if isinstance(m, ToolMessage) and m.name == PROFESSOR_TOOL_NAME:
                artifact = getattr(m, 'artifact', None)
                if artifact:
                    for d in artifact:
                        course = d.metadata.get('course', '')
                        prof = d.metadata.get('professor', '')
                        if course and course not in seen:
                            seen.add(course)
                            mentioned.append(f"{course} ({prof})")
        return mentioned, ''

    candidate_pairs = collect_searched_pairs(result)
    if not candidate_pairs:
        return [], ''

    answer_lower = answer.lower()
    mentioned_pairs = [pair for pair in candidate_pairs if pair[0] and pair[0] in answer]
    if not mentioned_pairs:
        mentioned_pairs = [pair for pair in candidate_pairs if pair[1] and pair[1].lower() in answer_lower]
    final_pairs = mentioned_pairs if mentioned_pairs else candidate_pairs

    seen_courses = set()
    deduped = []
    for course, prof in final_pairs:
        if course not in seen_courses:
            seen_courses.add(course)
            deduped.append((course, prof))

    return [f"{course} ({prof})" if prof else course for course, prof in deduped], ''

def used_general_knowledge(result): return any(isinstance(m, ToolMessage) and m.name == GENERAL_TOOL_NAME for m in result['messages'])
def used_web_search(result): return any(isinstance(m, ToolMessage) and m.name == WEB_SEARCH_TOOL_NAME for m in result['messages'])
def used_career_tool(result): return any(isinstance(m, ToolMessage) and m.name == CAREER_TOOL_NAME for m in result['messages'])
def used_professor_courses_tool(result): return any(isinstance(m, ToolMessage) and m.name == PROFESSOR_TOOL_NAME for m in result['messages'])
def used_any_tool(result): return any(isinstance(m, ToolMessage) for m in result['messages'])

def has_any_internal_data(result):
    for message in result['messages']:
        if isinstance(message, ToolMessage) and message.name == CAREER_TOOL_NAME:
            if getattr(message, 'artifact', None):
                return True
    return False

def build_header(n_reviews, not_found, is_general, used_web, is_career, course_names, professor_name, question_is_english=False, is_count=False):
    if question_is_english:
        if is_general: return '[AI General Answer · No Internal Data]'
        if is_count: return '[Based on Internal Data · Review Count]'
        if is_career:
            if not_found: return '[Not Enough Internal Data]'
            base = '[Internal Data (Reviews + Syllabi)]'
            return '[Internal Data (Reviews + Syllabi) + External Info]' if used_web else base
        if n_reviews > 0 and not not_found:
            base = f'[Based on Internal Data · {n_reviews} review(s)'
            base += ' + External Info]' if used_web else ']' 
            if course_names:
                return f"{base}\n\n**Courses (Professor):** {', '.join(course_names)}"
            return base
        if used_web: return '[External Info (Web Search) · No Internal Data]'
        return '[Not Enough Reviews Available]'

    if is_general: return '[AI 일반 답변 · 내부 자료 없음]'
    if is_count: return '[내부 자료 기반 · 후기 집계]'
    if is_career:
        if not_found: return '[참고할 내부 자료가 부족합니다]'
        base = '[내부 자료(후기+강의계획서) 참고]'
        return '[내부 자료(후기+강의계획서) + 외부 정보 참고]' if used_web else base
    if n_reviews > 0 and not not_found:
        base = f'[내부 자료 기반 · {n_reviews}명의 후기 기반'
        base += ' + 외부 정보 참고]' if used_web else ']' 
        if course_names:
            return f"{base}\n\n**과목 (교수님):** {', '.join(course_names)}"
        return base
    if used_web: return '[외부 정보(웹검색) 기반 · 내부 자료 없음]'
    return '[참고할 후기가 부족합니다]'

# JSON 파싱용 항목 매핑
import json as _json_module, re as _re_module

LABEL_MAP_KO = {
    '과목명': '과목명',

    # 진로 추천용
    '추천이유': '추천 이유',
    '얻을수있는역량': '얻을 수 있는 역량',
    '수업방식': '수업 방식',
    '평가방식': '평가 방식',
    '주의할점': '주의할 점',

    # 일반 후기용
    '개인과제': '개인 과제',
    '팀프로젝트': '팀 프로젝트',
    '수업분위기': '수업 분위기',
    '꿀팁추천대상': '꿀팁/추천대상',
    '시험': '시험',
    '교수스타일': '교수 스타일',
    '기타': '기타',
}

LABEL_MAP_EN = {
    '과목명': 'Course',

    # career recommendation
    '추천이유': 'Why this course fits',
    '얻을수있는역량': 'Skills you can build',
    '수업방식': 'Class format',
    '평가방식': 'Grading / assessment',
    '주의할점': 'Things to keep in mind',

    # general review
    '개인과제': 'Individual assignments',
    '팀프로젝트': 'Team project',
    '수업분위기': 'Class atmosphere',
    '꿀팁추천대상': 'Tips / recommended for',
    '시험': 'Exams',
    '교수스타일': 'Professor style',
    '기타': 'Other notes',
}

# 기존 코드와의 호환용 별칭
LABEL_MAP = LABEL_MAP_KO

def format_single_course(obj: dict, english: bool = False) -> str:
    label_map = LABEL_MAP_EN if english else LABEL_MAP

    merged = {}

    for key, value in obj.items():
        if key == '과목명':
            continue
        if value is None:
            continue

        label = label_map.get(key, key)

        if isinstance(value, list):
            value = ' '.join(str(v).strip() for v in value if str(v).strip())
        else:
            value = str(value).strip()

        if not value:
            continue

        if label in merged:
            merged[label] = merged[label].rstrip() + ' ' + value
        else:
            merged[label] = value

    lines = []
    for label, value in merged.items():
        lines.append(f'**{label}:** {value}')

    return '\n\n'.join(lines)


def ensure_multi_course_json_array(answer, course_names, result):
    """
    여러 과목이 검색됐는데 LLM이 단일 JSON 객체로 답한 경우,
    과목별 JSON 배열로 다시 작성하게 강제한다.
    """
    if len(course_names) <= 1:
        return answer

    try:
        clean = _re_module.sub(r'```json|```', '', str(answer)).strip()
        parsed = _json_module.loads(clean)

        # 이미 과목별 배열이면 그대로 사용
        if isinstance(parsed, list):
            return answer

        # 여러 과목인데 dict 하나면 잘못된 형식이므로 재작성 요청
        if isinstance(parsed, dict):
            rewrite_request = HumanMessage(
                "방금 답변은 여러 과목을 하나로 합쳐서 잘못 작성했어. "
                "검색 결과에 나온 과목이 2개 이상이면 반드시 과목별 JSON 배열로 나눠야 해. "
                "각 객체에는 반드시 \"과목명\" 키를 넣고, 각 과목의 후기 내용만 해당 객체에 작성해. "
                "서로 다른 과목의 내용을 섞지 마. "
                "사용자가 한국어로 질문했다면 값(value)은 한국어로, 영어로 질문했다면 값(value)은 영어로 작성해. "
                "코드블록 없이 순수 JSON 배열만 출력해.\n\n"
                f"대상 과목 목록:\n{course_names}\n\n"
                f"잘못된 이전 답변:\n{answer}"
            )
            rewrite_result = agent.invoke({'messages': result['messages'] + [rewrite_request]})
            return extract_text(rewrite_result)

    except Exception:
        pass

    return answer

def parse_json_answer(answer, english: bool = False):
    """LLM이 반환한 JSON을 파싱해서 포맷된 마크다운으로 변환한다."""
    try:
        clean = _re_module.sub(r'```json|```', '', answer).strip()
        parsed = _json_module.loads(clean)
        if isinstance(parsed, list):
            sections = []
            for item in parsed:
                if not isinstance(item, dict):
                    continue
                course_name = item.get('과목명', '')
                content = format_single_course(item, english=english)
                if content:
                    if course_name:
                        sections.append(f'**[{course_name}]**\n\n{content}')
                    else:
                        sections.append(content)
            if sections:
                return '\n\n---\n\n'.join(sections)
        elif isinstance(parsed, dict):
            course_name = parsed.get('과목명', '')
            result = format_single_course(parsed, english=english)
            if result:
                if course_name:
                    return f'**[{course_name}]**\n\n{result}'
                return result
    except Exception:
        pass
    return answer

def flatten_general_answer(answer):
    """
    일반 답변이 JSON으로 오면 value만 자연스럽게 이어 붙이고,
    JSON이 아니면 원문을 그대로 반환한다.
    """
    try:
        clean = _re_module.sub(r'```json|```', '', str(answer)).strip()
        parsed = _json_module.loads(clean)

        if isinstance(parsed, dict):
            parts = []
            for v in parsed.values():
                if isinstance(v, str) and v.strip():
                    parts.append(v.strip())
            if parts:
                return ' '.join(parts)

        return answer

    except Exception:
        return answer# 3. 여러 과목

def is_general_func(result): return used_general_knowledge({'messages': result['messages']})

def ask(question):
    result = agent.invoke({'messages': [HumanMessage(question)]})
    answer = extract_text(result)

    if not used_any_tool(result):
        answer = parse_json_answer(answer, english=is_english(question))
        if is_general_func(result):
            answer = flatten_general_answer(answer)
        return {'question': question, 'answer': answer, 'n_reviews': 0, 'result': result, 'header': ''}

    n_reviews = extract_n_reviews(result)
    is_general = used_general_knowledge(result)
    used_web = used_web_search(result)
    is_career = used_career_tool(result)
    is_professor_courses = used_professor_courses_tool(result)
    question_is_english = is_english(question)
    is_comparison = any(
    isinstance(m, ToolMessage) and m.name == 'course_comparison_tool'
    for m in result['messages']
)
    is_count = any(isinstance(m, ToolMessage) and m.name == 'count_reviews_tool' for m in result['messages'])

    if is_professor_courses or is_count:
        not_found = False
    elif is_career:
        not_found = not has_any_internal_data(result)
    else:
        not_found = any(phrase in answer for phrase in [
            '확인해 주시겠어요', '확인 부탁드립니다', '아직 이 과목에 대한 후기가 없', '후기가 없어요'
        ])

    course_names, professor_name = ([], '') if (not_found or is_general or is_career or is_count) else get_header_info(result, answer)
    header = build_header(n_reviews, not_found, is_general, used_web, is_career, course_names, professor_name, question_is_english, is_count)

    # 1. 영어 질문 + 한국어 답변 → 영어로 번역 (JSON 파싱 전에)
    if question_is_english and not is_english(answer):
        translate_request = HumanMessage(
            f"Please rewrite your previous answer entirely in natural, fluent English. "
            f"Keep the same structure, items, and information. Do not add or remove any content:\n\n{answer}"
        )
        translate_result = agent.invoke({'messages': result['messages'] + [translate_request]})
        answer = extract_text(translate_result)

    # 2. 한국어 질문 + 영어 답변 → 한국어로 번역 (JSON 파싱 전에)
    if not question_is_english and is_english(answer):
        translate_request = HumanMessage(
            f"아래 JSON을 그대로 유지하되, 값(value)만 자연스러운 한국어 ~해요 말투로 번역해줘. "
            f"키(key)는 절대 바꾸지 마. JSON 형식 그대로 출력해:\n\n{answer}"
        )
        translate_result = agent.invoke({'messages': result['messages'] + [translate_request]})
        answer = extract_text(translate_result)

    
    # 3. 비교 질문이면 JSON 포맷으로 보내지 말고, Markdown 비교표로 재정리
    if is_comparison:
        target_language = "English" if question_is_english else "Korean"

        comparison_rewrite_prompt = HumanMessage(
            f"""
    You are rewriting a course comparison answer for an MBA course registration chatbot.

    The user's question was:
{question}

    The draft answer was:
{answer}

    Rewrite the answer as a clear Markdown comparison table.

    Rules:
    - Write in {target_language}.
    - Do not repeat the same content twice.
    - Do not add information that is not in the draft answer.
    - Use the exact course names that appear in the draft.
    - Make it useful for course registration decisions.
    - If information is missing, write that it is not clearly found in the available reviews.
    - Do not use JSON.
    - Do not use code blocks.

    Required structure:

    ## 과목 비교 / Course Comparison: Course A vs Course B

    ### 1. 한눈에 비교 / Quick Comparison

    | 항목 / Category | Course A | Course B |
    |---|---|---|
    | 수업방식 / Class format | ... | ... |
    | 과제·팀플 부담 / Assignments & team projects | ... | ... |
    | 시험 부담 / Exam difficulty | ... | ... |
    | 수업 분위기 / Class atmosphere | ... | ... |
    | 실무성 / Practical value | ... | ... |
    | 주의할 점 / Things to keep in mind | ... | ... |

    ### 2. 과목별 핵심 요약 / Course-by-course summary

    #### Course A
    ...

    #### Course B
    ...

    ### 3. 최종 추천 / Final recommendation
    ...
    """
    )

        rewrite_response = model.invoke([comparison_rewrite_prompt])
        answer = extract_message_text(rewrite_response)

    # 4. 일반 과목/교수님 후기 질문이면 기존 JSON 파싱/포맷팅
    elif not (is_general or is_count or not_found):
        answer = ensure_multi_course_json_array(answer, course_names, result)
        answer = parse_json_answer(answer, english=question_is_english)

    # 5. 후기 집계 결과를 테이블로 변환
    if is_count:
        try:
            import json as _j
            _parsed = _j.loads(answer)
            if isinstance(_parsed, list):
                lines = ['| 순위 | 과목명 | 후기 수 |', '|---|---|---|']
                for i, item in enumerate(_parsed, 1):
                    review_count = str(item.get('후기수', '')).replace('건', '').strip()
                    lines.append(f"| {i} | {item.get('과목명','')} | {review_count}건 |")
                answer = '\n'.join(lines)
        except Exception:
            pass

    # 6. 일반 답변이 JSON으로 오면 {} 없이 value만 이어 붙임
    if is_general:
        answer = flatten_general_answer(answer)

# ── 사이드바 ──
with st.sidebar:
    st.markdown('## ☕ 라떼는 말이야')
    st.markdown('SNU MBA 후배를 위한 AI 선배 챗봇')
    st.divider()
    st.markdown('**이런 질문을 해봐요:**')
    st.markdown('※ 과목명은 한국어/영어 그대로 입력해주세요  \nPlease use the exact course name in Korean or English')
    for ex in [
        '박성호 교수님 Marketing Analytics 수업 어때?',
        '박진수 교수님 수업 뭐 있어?',
        'Portfolio Management랑 Financial Engineering 비교해줘',
        'OO 회사 / OO 직무에 관심 있는데, 도움 되는 수업 추천해줘',
        '후기 많은 과목 순서대로 보여줘',
        '수강후기 남길게요',
    ]:
        if st.button(ex, key=ex):
            st.session_state['suggested'] = ex
    st.divider()
    if st.button('🗑️ 대화 초기화'):
        st.session_state.update({'messages': [], 'conversation_history': [], 'is_first_turn': True})
        st.rerun()
    st.markdown('---')
    st.markdown('채팅창에 `수강후기 남길게요` 입력!')


# ── 메인 ──
st.title('☕ 라떼는 말이야')
st.markdown('**SNU MBA 선배들의 솔직한 수강후기 — AI가 정리해드려요**')
st.markdown('※ 과목명은 한국어/영어 그대로 입력해주시면 더 정확해요 | Please use the exact course name for best results')

if st.session_state.get('review_success_message'):
    st.success(st.session_state.pop('review_success_message'))

try:
    resources = load_all()
    all_docs = resources['all_docs']
    review_docs = resources['review_docs']
    review_vectorstore = resources['review_vectorstore']
    vectorstore = resources['vectorstore']
    professor_aliases = resources['professor_aliases']
    review_sheet = resources['review_sheet']
    agent = resources['agent']
except Exception as e:
    st.error(f'데이터 로드 실패: {e}')
    st.stop()

for msg in st.session_state['messages']:
    with st.chat_message(msg['role']):
        st.markdown(msg['content'])

prompt = st.session_state.pop('suggested', None)

# ── 후기 입력 폼 ──
if st.session_state.get('review_mode'):
    with st.form('review_form'):
        st.markdown('### 📝 수강후기 입력')
        review_text = st.text_area(
            '과목명, 교수님, 후기 내용을 자유롭게 입력해주세요!',
            placeholder='예: 박진수 교수님 AI 수업 들었는데 AI를 이해하는데 도움이 되었습니다'
        )
        author = st.text_input('이름/닉네임 (선택)', placeholder='익명')
        col_submit, col_cancel = st.columns(2)
        with col_submit:
            submitted = st.form_submit_button('후기 저장하기 ☕')
        with col_cancel:
            cancelled = st.form_submit_button('취소')

        if cancelled:
            st.session_state['review_mode'] = False
            st.rerun()

        if submitted and review_text.strip():
            with st.spinner('🤖 분석 중...'):
                try:
                    from langchain.chat_models import init_chat_model as _icm2
                    from langchain_core.documents import Document as _LCDoc
                    from langchain.messages import HumanMessage as _HM2
                    import json as _jj, re as _rr
                    _tag_model = _icm2(model='openai:gpt-4o-mini')
                    known_courses = list({d.metadata.get('course', '') for d in all_docs if d.metadata.get('type') == 'review'})
                    known_profs = list({d.metadata.get('professor', '') for d in all_docs if d.metadata.get('type') == 'review'})
                    prompt_tag = f"""아래 후기에서 정보를 추출해서 JSON으로만 답해라.
알려진 과목명: {known_courses}
알려진 교수명: {known_profs}
추출: course(과목명), professor(교수명, 모르면 ""), review(후기 내용만)
텍스트: {review_text}
JSON:"""
                    resp = _tag_model.invoke([_HM2(prompt_tag)])
                    text_resp = _rr.sub(r'```json|```', '', resp.content.strip()).strip()
                    info = _jj.loads(text_resp)
                    course = info.get('course', '').strip()
                    professor = info.get('professor', '').strip()
                    review_content = info.get('review', review_text).strip()
                    author_name = author.strip() or '익명'
                    input_date = datetime.now().strftime('%Y-%m-%d')
                    language = ''
                    existing = [d for d in all_docs if d.metadata.get('course', '').lower() == course.lower() and d.metadata.get('type') == 'review']
                    if existing:
                        language = existing[0].metadata.get('language', '')
                    alt_name = professor_aliases.get(professor, '')
                    prof_display = f"{professor} ({alt_name})" if alt_name else professor
                    new_doc = _LCDoc(
                        page_content=(f'[과목명: {course}] [교수: {prof_display}] [언어: {language}] [수강시기: {input_date}]\n후기: {review_content}'),
                        metadata={'type': 'review', 'course': course, 'professor': professor,
                                  'language': language, 'term': input_date, 'author': author_name}
                    )
                    review_vectorstore.add_documents([new_doc])
                    all_docs.append(new_doc)
                    review_docs.append(new_doc)
                    if review_sheet:
                        review_sheet.append_row([author_name, course, professor, language, input_date, review_content])

                    success_message = f'☕ 후기 감사해요! [{course}] 과목 후기가 저장됐어요!'

                    st.session_state['review_success_message'] = success_message
                    st.session_state['messages'].append({
                        'role': 'assistant',
                        'content': success_message
                    })
                    st.session_state['review_mode'] = False
                    st.rerun()
                except Exception as e:
                    st.error(f'저장 실패: {e}')

typed_input = st.chat_input('궁금한 걸 입력하세요 / Ask anything')
user_input = prompt or typed_input

if user_input:
    with st.chat_message('user'):
        st.markdown(user_input)
    st.session_state['messages'].append({'role': 'user', 'content': user_input})

    with st.chat_message('assistant'):
        with st.spinner('☕ 선배들의 후기를 찾아보는 중...'):
            try:
                from langchain.messages import HumanMessage
                from langchain_core.messages import ToolMessage
                review_keywords = [
                    '수강후기', '후기 남', '후기를 남', '후기 입력', '후기남길',
                    'leave a review', 'write a review', 'submit a review',
                    'add a review', 'leave review', 'write review',
                ]
                if any(kw in user_input.lower() for kw in review_keywords):
                    st.session_state['review_mode'] = True
                    full_answer = '📝 후기 입력 모드예요! 아래 폼에 입력해주세요.'
                    st.session_state['messages'].append({'role': 'assistant', 'content': full_answer})
                    st.rerun()
                else:
                    conversation_history = st.session_state['conversation_history']
                    turn_start = len(conversation_history)
                    conversation_history.append(HumanMessage(user_input))

                    result = agent.invoke({'messages': conversation_history})
                    conversation_history = result['messages']
                    st.session_state['conversation_history'] = conversation_history

                    new_messages = conversation_history[turn_start:]
                    turn_result = {'messages': new_messages}

                    answer = extract_text(result)
                    question_is_english = is_english(user_input)

                    is_general = used_general_knowledge(turn_result)
                    used_web = used_web_search(turn_result)
                    is_career = used_career_tool(turn_result)
                    is_professor_courses = used_professor_courses_tool(turn_result)
                    is_count = any(isinstance(m, ToolMessage) and m.name == 'count_reviews_tool' for m in new_messages)
                    is_comparison = any(isinstance(m, ToolMessage) and m.name == 'course_comparison_tool' for m in new_messages)
                    n_reviews = extract_n_reviews(turn_result)

                    if not used_any_tool(turn_result):
                        answer = parse_json_answer(answer, english=question_is_english)
                        if is_general:
                            answer = flatten_general_answer(answer)
                        st.markdown(answer)
                        st.session_state['messages'].append({'role': 'assistant', 'content': answer})
                    else:
                        if is_professor_courses or is_count:
                            not_found = False
                        elif is_career:
                            not_found = not has_any_internal_data(turn_result)
                        else:
                            not_found = any(phrase in answer for phrase in [
                                '확인해 주시겠어요', '확인 부탁드립니다', '아직 이 과목에 대한 후기가 없', '후기가 없어요'
                            ])

                        course_names, professor_name = ([], '') if (not_found or is_general or is_career or is_count) else get_header_info(turn_result, answer)
                        header = build_header(n_reviews, not_found, is_general, used_web, is_career, course_names, professor_name, question_is_english, is_count)

                        if question_is_english and not is_english(answer):
                            translate_request = HumanMessage(
                                f"Please rewrite your previous answer entirely in natural, fluent English. "
                                f"Keep the same structure, items, and information. Do not add or remove any content:\n\n{answer}"
                            )
                            translate_result = agent.invoke({'messages': conversation_history + [translate_request]})
                            answer = extract_text(translate_result)

                        if not question_is_english and is_english(answer):
                            translate_request = HumanMessage(
                                f"아래 JSON을 그대로 유지하되, 값(value)만 자연스러운 한국어 ~해요 말투로 번역해줘. "
                                f"키(key)는 절대 바꾸지 마. JSON 형식 그대로 출력해:\n\n{answer}"
                            )
                            translate_result = agent.invoke({'messages': conversation_history + [translate_request]})
                            answer = extract_text(translate_result)

                        if is_comparison:
                            from langchain.chat_models import init_chat_model as _icm
                            _model = _icm(model='openai:gpt-4.1-mini')
                            target_language = "English" if question_is_english else "Korean"
                            comparison_sources = get_comparison_sources_text(turn_result)
                            comp_prompt = HumanMessage(
                                f"You are an MBA course registration assistant.\n"
                                f"User asked: {user_input}\n"
                                f"Sources: {comparison_sources}\n"
                                f"Write a clear Markdown comparison in {target_language}. No JSON, no code blocks.\n"
                                f"Structure: ## Course Comparison, ### 1. Quick Comparison (table), ### 2. Summaries, ### 3. Recommendation"
                            )
                            comp_resp = _model.invoke([comp_prompt])
                            answer = comp_resp.content if isinstance(comp_resp.content, str) else str(comp_resp.content)

                        elif not (is_general or is_count or not_found):
                            answer = ensure_multi_course_json_array(answer, course_names, result)
                            answer = parse_json_answer(answer, english=question_is_english)

                        if is_count:
                            try:
                                import json as _j
                                _parsed = _j.loads(answer)
                                if isinstance(_parsed, list):
                                    lines = ['| 순위 | 과목명 | 후기 수 |', '|---|---|---|']
                                    for i, item in enumerate(_parsed, 1):
                                        review_count = str(item.get('후기수', '')).replace('건', '').strip()
                                        lines.append(f"| {i} | {item.get('과목명','')} | {review_count}건 |")
                                    answer = '\n'.join(lines)
                            except Exception:
                                pass

                        if is_general:
                            answer = flatten_general_answer(answer)

                        exam_questions = extract_exam_questions_from_result(turn_result)

                        if exam_questions:
                            answer += f'\n\n**시험기출문제:**\n{exam_questions}'

                        if st.session_state['is_first_turn'] and not is_general and not is_career and not is_count and not is_comparison:
                            if question_is_english:
                                answer += (
                                    '\n\n---\n'
                                    '💬 Have you taken any courses recently? Could you leave a quick one-line review? '
                                    'If you type "leave a review", I\'ll help you right away!\n\n'
                                    '💬 혹시 최근에 들으신 수업이 있다면, 후기를 한 줄만 남겨주실 수 있어요? '
                                    '"수강후기 남길게요"라고 입력하시면 바로 도와드릴게요!'
                                )
                            else:
                                answer += (
                                    '\n\n---\n'
                                    '💬 혹시 최근에 들으신 수업이 있다면, 후기를 한 줄만 남겨주실 수 있어요? '
                                    '"수강후기 남길게요"라고 입력하시면 바로 도와드릴게요!\n\n'
                                    '💬 Have you taken any courses recently? Could you leave a quick one-line review? '
                                    'If you type "leave a review", I\'ll help you right away!'
                                )
                            st.session_state['is_first_turn'] = False

                        full_answer = f"{header}\n\n{answer}" if header else answer
                        st.markdown(full_answer)
                        st.session_state['messages'].append({'role': 'assistant', 'content': full_answer})

            except Exception as e:
                st.error(f'오류: {e}')

st.divider()
col1, col2, col3 = st.columns(3)
with col1:
    st.metric('📝 총 후기 수', f"{sum(1 for d in all_docs if d.metadata.get('type')=='review')}건")
with col2:
    st.metric('📚 과목 수', f"{len({d.metadata.get('course') for d in all_docs if d.metadata.get('type')=='review'})}개")
with col3:
    st.metric('☁️ Google Sheets', '✅ 연결됨' if review_sheet else '⚠️ 미연결')
